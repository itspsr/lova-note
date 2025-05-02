import os
import whisper
import librosa
import torch
import soundfile as sf
import torchaudio
import yt_dlp
import re
from datetime import datetime
from fpdf import FPDF
from docx import Document
import openai
from dotenv import load_dotenv

# === Constants ===
UPLOAD_FOLDER = 'uploads'
TRANSCRIPT_FOLDER = 'transcriptions'

# === Load environment variables ===
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")
if not openai.api_key:
    raise ValueError("OpenAI API key is not set. Please set the OPENAI_API_KEY environment variable.")

LANGUAGE_NAMES = {
    "en": "English", "hi": "Hindi", "bn": "Bengali", "te": "Telugu", "ta": "Tamil",
    "gu": "Gujarati", "ml": "Malayalam", "kn": "Kannada", "mr": "Marathi", "pa": "Punjabi", "ur": "Urdu"
}

def get_language_name(code):
    return LANGUAGE_NAMES.get(code, code)

def get_audio_duration(filepath):
    audio, sr = librosa.load(filepath, sr=None)
    return librosa.get_duration(y=audio, sr=sr)

def detect_language(filepath, model_size='base'):
    model = whisper.load_model(model_size)
    audio = whisper.load_audio(filepath)
    audio = whisper.pad_or_trim(audio)
    mel = whisper.log_mel_spectrogram(audio).to(model.device)

    _, probs = model.detect_language(mel)
    lang_probs = {k: float(v) for k, v in probs.items()}
    detected_lang = max(lang_probs, key=lang_probs.get)
    confidence = lang_probs[detected_lang]

    if detected_lang not in LANGUAGE_NAMES:
        print(f"[Warning] Detected unknown language '{detected_lang}', defaulting to English.")
        detected_lang = "en"

    return str(detected_lang), confidence

def preprocess_audio(filepath):
    waveform, sample_rate = torchaudio.load(filepath)
    if waveform.shape[0] > 1:
        waveform = torch.mean(waveform, dim=0, keepdim=True)
    if sample_rate != 16000:
        resampler = torchaudio.transforms.Resample(orig_freq=sample_rate, new_freq=16000)
        waveform = resampler(waveform)
    waveform = waveform / waveform.abs().max()
    base, _ = os.path.splitext(filepath)
    temp_path = f"{base}_preprocessed.wav"
    sf.write(temp_path, waveform.squeeze(0).numpy(), 16000)
    return temp_path

def clean_transcription(text):
    cleaned_text = re.sub(r'\b(um|uh|ah|like|you know)\b', '', text)
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that cleans and improves transcription text."},
                {"role": "user", "content": f"Clean and perfect the following transcription text for grammar, punctuation, and readability:\n{text}"}
            ],
            temperature=0.2,
            max_tokens=1500
        )
        cleaned_text = response.choices[0].message['content'].strip()
    except Exception as e:
        print(f"[GPT CLEANING ERROR] {e}")
        cleaned_text = text
    return cleaned_text

def transcribe_audio(filepath, model_size='base', lang='auto'):
    preprocessed_path = preprocess_audio(filepath)
    model = whisper.load_model(model_size)

    if lang == 'auto':
        lang, confidence = detect_language(preprocessed_path, model_size)
    else:
        confidence = None

    result = model.transcribe(preprocessed_path, language=lang)
    cleaned_text = clean_transcription(result['text'])

    return {
        'text': cleaned_text,
        'language': lang,
        'confidence': confidence,
        'accuracy_rate': (confidence * 100) if confidence is not None else 100.0,
        'duration': get_audio_duration(preprocessed_path),
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'audio_path': preprocessed_path
    }

def export_to_pdf(text, metadata, filename="transcription.pdf"):
    path = os.path.join(TRANSCRIPT_FOLDER, filename)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)

    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "Transcription Summary", ln=True)
    pdf.set_font("Arial", size=12)
    pdf.cell(0, 10, f"Language: {get_language_name(metadata['language'])}", ln=True)
    if metadata.get("accuracy_rate") is not None:
        pdf.cell(0, 10, f"Confidence: {metadata['accuracy_rate']:.2f}%", ln=True)
    pdf.cell(0, 10, f"Duration: {metadata['duration']:.2f} seconds", ln=True)
    pdf.cell(0, 10, f"Timestamp: {metadata['timestamp']}", ln=True)
    pdf.cell(0, 10, f"Audio File Path: {metadata['audio_path']}", ln=True)

    pdf.ln(10)
    pdf.multi_cell(0, 10, text)
    pdf.output(path)

    return filename

def export_to_docx(text, metadata, filename="transcription.docx"):
    path = os.path.join(TRANSCRIPT_FOLDER, filename)
    doc = Document()
    doc.add_heading('Transcription Summary', 0)

    doc.add_paragraph(f"Language: {get_language_name(metadata['language'])}")
    if metadata.get("accuracy_rate") is not None:
        doc.add_paragraph(f"Confidence: {metadata['accuracy_rate']:.2f}%")
    doc.add_paragraph(f"Duration: {metadata['duration']:.2f} seconds")
    doc.add_paragraph(f"Timestamp: {metadata['timestamp']}")
    doc.add_paragraph(f"Audio File Path: {metadata['audio_path']}")

    doc.add_paragraph()
    doc.add_paragraph(text)
    doc.save(path)

    return filename

def download_youtube_audio(url):
    ydl_opts = {
        'format': 'bestaudio/best',
        'outtmpl': os.path.join(UPLOAD_FOLDER, '%(id)s.%(ext)s'),
        'postprocessors': [{
            'key': 'FFmpegAudioConvertor',
            'preferredcodec': 'mp3',
            'preferredquality': '192',
        }],
    }
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        result = ydl.extract_info(url, download=True)
        return os.path.join(UPLOAD_FOLDER, f"{result['id']}.mp3")
